"""Parser selection and structured-block extraction."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

from app.parsers.errors import PermanentParseError, UnsupportedMimeError
from app.parsers.models import BlockType, RawDocument
from app.parsers.ocr import should_ocr
from app.parsers.registry import ParserRegistry, get_parser_registry


def _raw(mime: str, data: bytes, filename: str = "doc") -> RawDocument:
    return RawDocument(
        data=data,
        mime_type=mime,
        filename=filename,
        title=Path(filename).stem,
    )


def test_registry_selects_parser_by_mime() -> None:
    registry = get_parser_registry()
    assert registry.select("text/plain").name == "txt"
    assert registry.select("text/markdown").name == "markdown"
    assert registry.select("text/html").name == "html"
    assert registry.select("application/pdf").name == "pdf"
    assert (
        registry.select(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ).name
        == "docx"
    )
    with pytest.raises(UnsupportedMimeError):
        registry.select("image/png")


def test_txt_parser_emits_structured_blocks() -> None:
    source = b"""# Leave Policy

Employees receive 22 days annual leave.

- Public holidays are extra
- Unused days carry over

> Check with HR before booking.

    code_example = True
"""
    parsed = get_parser_registry().parse(_raw("text/plain", source, "handbook.txt"))
    types = [block.type for block in parsed.blocks]
    assert BlockType.TITLE in types
    assert BlockType.HEADING in types or BlockType.TITLE in types
    assert BlockType.PARAGRAPH in types
    assert BlockType.LIST in types
    assert BlockType.QUOTE in types
    assert BlockType.CODE in types
    assert parsed.parser_name == "txt"
    assert parsed.parser_version == 1
    assert parsed.blocks[0].type is BlockType.TITLE


def test_markdown_parser_preserves_headings_lists_code_and_tables() -> None:
    source = b"""# Authentication

## Tokens

Use a **bearer** token.

- Access token
- Refresh token

```python
print("ok")
```

| Method | Path |
| --- | --- |
| GET | /health |

![Diagram](auth.png "Login flow")
"""
    parsed = get_parser_registry().parse(_raw("text/markdown", source, "auth.md"))
    types = {block.type for block in parsed.blocks}
    assert parsed.title == "Authentication"
    assert BlockType.HEADING in types
    assert BlockType.LIST in types
    assert BlockType.CODE in types
    assert BlockType.TABLE in types
    assert BlockType.IMAGE_CAPTION in types
    heading = next(block for block in parsed.blocks if block.type is BlockType.HEADING)
    assert heading.level == 2
    assert heading.text == "Tokens"
    paragraph = next(block for block in parsed.blocks if block.type is BlockType.PARAGRAPH)
    assert "bearer" in paragraph.text
    assert "**" not in paragraph.text
    table = next(block for block in parsed.blocks if block.type is BlockType.TABLE)
    assert table.metadata["rows"][0] == ["Method", "Path"]
    code = next(block for block in parsed.blocks if block.type is BlockType.CODE)
    assert "print" in code.text
    assert code.metadata.get("language") == "python"


def test_html_parser_strips_scripts_and_walks_structure() -> None:
    source = b"""<html><head><title>Runbook</title>
<script>alert(1)</script><style>body{}</style></head>
<body>
<h1>Runbook</h1>
<p>Restart Redis when memory is high.</p>
<ol><li>Flush cache</li><li>Restart</li></ol>
<table><tr><th>Error</th><th>Action</th></tr>
<tr><td>OOM</td><td>Restart</td></tr></table>
<blockquote>Page on-call if restart fails.</blockquote>
<img alt="Redis dashboard" src="dash.png"/>
</body></html>"""
    parsed = get_parser_registry().parse(_raw("text/html", source, "runbook.html"))
    texts = " ".join(block.text for block in parsed.blocks)
    assert "alert" not in texts
    assert "body{}" not in texts
    types = {block.type for block in parsed.blocks}
    assert BlockType.PARAGRAPH in types
    assert BlockType.LIST in types
    assert BlockType.TABLE in types
    assert BlockType.QUOTE in types
    assert BlockType.IMAGE_CAPTION in types
    assert parsed.title == "Runbook"


def test_docx_parser_reads_headings_and_tables() -> None:
    document = DocxDocument()
    document.add_heading("Payment Service", level=1)
    document.add_heading("Refunds", level=2)
    document.add_paragraph("Refunds post to the ledger within two hours.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Type"
    table.cell(1, 0).text = "amount"
    table.cell(1, 1).text = "integer"
    buffer = BytesIO()
    document.save(buffer)
    parsed = get_parser_registry().parse(
        _raw(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
            "payments.docx",
        )
    )
    types = {block.type for block in parsed.blocks}
    assert parsed.title == "Payment Service"
    assert BlockType.HEADING in types
    assert BlockType.PARAGRAPH in types
    assert BlockType.TABLE in types
    table_block = next(block for block in parsed.blocks if block.type is BlockType.TABLE)
    assert "amount" in table_block.text


def test_pdf_parser_extracts_digital_text_without_ocr() -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=18)
    pdf.cell(0, 12, "Leave Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Employees receive 22 days annual leave each calendar year.")
    data = bytes(pdf.output())
    parsed = get_parser_registry().parse(_raw("application/pdf", data, "leave.pdf"))
    assert parsed.used_ocr is False
    assert parsed.page_count == 1
    blob = " ".join(block.text for block in parsed.blocks)
    assert "22 days" in blob
    assert any(block.page == 1 for block in parsed.blocks if block.type is not BlockType.TITLE)


def test_should_ocr_only_when_text_is_sparse() -> None:
    assert should_ocr(0, 2, min_chars_per_page=40) is True
    assert should_ocr(5000, 2, min_chars_per_page=40) is False


def test_corrupt_pdf_is_permanent() -> None:
    with pytest.raises(PermanentParseError):
        get_parser_registry().parse(_raw("application/pdf", b"%PDF-not-really", "bad.pdf"))


def test_parser_identity_is_versioned() -> None:
    name, version = ParserRegistry().identity("text/markdown")
    assert name == "markdown"
    assert version == 1
